import docx
import openai
import os

from fastapi import FastAPI, File, UploadFile
from fastapi.responses import FileResponse
from googletrans import Translator
from dotenv import load_dotenv

load_dotenv()

# Initialize the OpenAI API with your API key
openai.api_key = os.environ.get('OPENAI_SECRET_KEY')

# Initialize the FastAPI app
app = FastAPI()

translate_client = Translator()

# Define function to rephrase a sentence using OpenAI GPT
def rephrase_sentence(sentence):
    rephrased_sentence = openai.Completion.create(
        engine="text-davinci-002",
        prompt=f"Rephrase the following text: '{sentence}'",
        temperature=0.5,
        max_tokens=400,
        n=1,
        stop=None,
        timeout=20,
    ).choices[0].text.strip()
    return rephrased_sentence

# Define function to translate a sentence using Google Translate API
def translate_sentence(sentence, dest_language):
    translated_text = translate_client.translate(sentence, dest=dest_language).text
    return translated_text

def create_translated_table(reprhased_rows, dest_language):
    # Create a new Word document
    translated_document = docx.Document()
    
    # Initialize a new table to store the translated rows
    translated_table = translated_document.add_table(rows=0, cols=2)
    translated_table.style = "Table Grid"
    
    # Translate each row in the table
    for row in reprhased_rows:
        # Get the English text from the first cell in the row
        original_en_text = row[0]
        rephrased_en_text = row[1]

        # Use the Google Translate API to translate the text to Arabic
        original_translation = translate_client.translate(original_en_text, dest=dest_language).text
        rephrased_translation = translate_client.translate(rephrased_en_text, dest=dest_language).text
        
        original_ar_text = original_translation
        rephrased_ar_text = rephrased_translation    
        
        original_new_row = translated_table.add_row()
        original_new_row.cells[0].text = original_en_text
        original_new_row.cells[1].text = original_ar_text
        
        rephrased_new_row = translated_table.add_row()
        rephrased_new_row.cells[0].text = rephrased_en_text
        rephrased_new_row.cells[1].text = rephrased_ar_text
    return translated_document
        
# Define the endpoint to translate the Word document
@app.post("/translate/")
async def translate_document(file: UploadFile = File(...)):

    # Save the uploaded file to the "uploads" directory
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as file_object:
        file_object.write(await file.read())

    # Load the Word document using the python-docx library
    doc = docx.Document(file_location)

    # Initialize an empty list to store the translated rows
    reprhased_rows = []

    # Loop over each row in the table and rephrase the sentence using the OpenAI GPT model
    for table in doc.tables:
        for row in table.rows:
            # Extract the text from the first cell of the row
            sentence = row.cells[0].text.strip()
            
            # Use the OpenAI GPT model to rephrase the sentence
            rephrased_sentence = rephrase_sentence(sentence)
            
            # Append the rephrased sentence to the list of translated rows
            reprhased_rows.append([sentence, rephrased_sentence])

    
    translated_document = create_translated_table(reprhased_rows, 'ar')
    # Save the translated document to a new file
    translated_file_path = f"uploads/{file.filename.split('.')[0]}_translated.docx"
    translated_document.save(translated_file_path)

    # Return the translated Word document to the client as a file download
    return FileResponse(
        path=translated_file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=translated_file_path
    )